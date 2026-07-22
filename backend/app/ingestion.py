"""
Document ingestion pipeline: extract text → chunk → embed → extract entities → store.

Supports PDF, DOCX, XLSX, and plain text files.
"""
import io
import logging
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models import Document, Chunk, Entity, Relationship
from app import gemini_client

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from an XLSX file — concatenates all cells across sheets."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
    text_parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        text_parts.append(f"=== Sheet: {sheet} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text_via_ocr(file_bytes: bytes) -> str:
    """Run OCR on image bytes (PNG, JPG) using pytesseract."""
    import pytesseract
    from PIL import Image
    try:
        logger.info("Running image OCR via pytesseract...")
        image = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(image)
    except Exception as e:
        logger.error(f"Image OCR extraction failed: {e}")
        return ""


def extract_text_from_pdf_ocr(file_bytes: bytes) -> str:
    """Fallback to convert PDF pages to images and run OCR using pdf2image and pytesseract."""
    from pdf2image import convert_from_bytes
    import pytesseract
    try:
        logger.info("PDF has no standard text. Running OCR fallback...")
        images = convert_from_bytes(file_bytes)
        text_parts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img)
            text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF OCR extraction failed: {e}")
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch text extraction based on file extension.
    Falls back to image OCR for PNG/JPG/JPEG, and triggers OCR fallback for empty PDFs.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "xlsx": extract_text_from_xlsx,
        "png": extract_text_via_ocr,
        "jpg": extract_text_via_ocr,
        "jpeg": extract_text_via_ocr,
    }

    text = ""
    if ext in extractors:
        text = extractors[ext](file_bytes)
    else:
        # Plain text fallback
        text = file_bytes.decode("utf-8", errors="replace")

    # Trigger scanned PDF OCR fallback if extracted text is empty or very short
    if ext == "pdf" and len(text.strip()) < 100:
        text = extract_text_from_pdf_ocr(file_bytes)

    return text


# ---------------------------------------------------------------------------
# Text Chunking
# ---------------------------------------------------------------------------


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Split text into chunks of approximately `chunk_size` words with `overlap` word overlap.
    Uses word-level splitting as a proxy for token count.
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]
        chunk = " ".join(chunk_words)
        if chunk.strip():
            chunks.append(chunk)
        # Advance by (chunk_size - overlap) words
        start += chunk_size - overlap

    return chunks


# ---------------------------------------------------------------------------
# Full Ingestion Pipeline
# ---------------------------------------------------------------------------


async def ingest_document(
    file_bytes: bytes,
    filename: str,
    db: AsyncSession,
) -> Document:
    """
    Full ingestion pipeline:
    1. Extract raw text from the file
    2. Chunk the text (~500 words per chunk with overlap)
    3. Generate embeddings via Gemini for each chunk
    4. Extract structured entities via Gemini
    5. Store everything in the database
    
    Returns the created Document record.
    """
    logger.info(f"Starting ingestion for: {filename}")

    # Step 1: Extract text
    raw_text = extract_text(file_bytes, filename)
    if not raw_text.strip():
        raise ValueError(f"No text could be extracted from {filename}")

    logger.info(f"Extracted {len(raw_text)} chars from {filename}")

    # Step 2: Chunk the text
    chunks = chunk_text(raw_text)
    logger.info(f"Created {len(chunks)} chunks from {filename}")

    # Step 3: Generate embeddings (batch)
    embeddings = []
    # Process in batches of 20 to avoid API limits
    batch_size = 20
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        batch_embeddings = await gemini_client.generate_embeddings_batch(batch)
        embeddings.extend(batch_embeddings)
    
    logger.info(f"Generated {len(embeddings)} embeddings for {filename}")

    # Step 4: Extract entities via Gemini
    entity_data = await gemini_client.extract_entities(raw_text)
    logger.info(f"Extracted entities from {filename}: {list(entity_data.keys())}")

    # Step 5: Store in database
    doc = Document(
        filename=filename,
        content_text=raw_text,
        doc_type=entity_data.get("document_type", "other"),
        metadata_json=entity_data,
        chunk_count=len(chunks),
    )
    db.add(doc)
    await db.flush()  # get the doc.id

    # Store chunks with embeddings
    for idx, (chunk_text_content, embedding) in enumerate(zip(chunks, embeddings)):
        chunk = Chunk(
            document_id=doc.id,
            content=chunk_text_content,
            chunk_index=idx,
            embedding=embedding,
        )
        db.add(chunk)

    # Store entities
    entity_type_map = {
        "equipment_tags": "equipment_tag",
        "dates": "date",
        "personnel": "person",
        "regulations": "regulation",
        "key_findings": "finding",
    }

    inserted_entities = []
    for json_key, entity_type in entity_type_map.items():
        values = entity_data.get(json_key, [])
        for value in values:
            val_str = str(value).strip()
            if val_str:
                entity = Entity(
                    document_id=doc.id,
                    entity_type=entity_type,
                    entity_value=val_str,
                    raw_context=entity_data.get("summary", ""),
                )
                db.add(entity)
                await db.flush()

                inserted_entities.append({
                    "id": entity.id,
                    "type": entity_type,
                    "value": val_str
                })

                # Create relationship linking entity to document
                rel = Relationship(
                    entity_id=entity.id,
                    document_id=doc.id,
                    relationship_type="extracted_from",
                )
                db.add(rel)

    # 1. Knowledge Graph: Extract and store entity-to-entity relationships
    if inserted_entities:
        try:
            relationships_data = await gemini_client.extract_entity_relationships(raw_text, inserted_entities)
            
            # Map lowercase entity values to their IDs
            value_to_entity = {e["value"].lower(): e["id"] for e in inserted_entities}
            
            from app.models import EntityRelationship
            for rel_data in relationships_data:
                src_val = str(rel_data.get("source_value", "")).strip().lower()
                tgt_val = str(rel_data.get("target_value", "")).strip().lower()
                rel_type = str(rel_data.get("relationship_type", "associated_with")).strip()
                
                if src_val in value_to_entity and tgt_val in value_to_entity:
                    entity_rel = EntityRelationship(
                        source_id=value_to_entity[src_val],
                        target_id=value_to_entity[tgt_val],
                        relationship_type=rel_type,
                        document_id=doc.id
                    )
                    db.add(entity_rel)
            logger.info(f"Extracted {len(relationships_data)} entity-to-entity relationships for {filename}")
        except Exception as e:
            logger.error(f"Failed to extract relationships for {filename}: {e}")

    # 2. Compliance Intelligence: Audit the document if relevant
    doc_type = entity_data.get("document_type", "other")
    compliance_types = {"safety_inspection", "regulatory_checklist", "incident_report", "maintenance_procedure"}
    if doc_type in compliance_types or any(x in filename.lower() for x in ["compliance", "safety", "audit", "checklist", "incident"]):
        try:
            audit_res = await gemini_client.audit_compliance(raw_text)
            from app.models import ComplianceFinding
            compliance_finding = ComplianceFinding(
                document_id=doc.id,
                regulation_type=audit_res.get("regulation_type", "General Safety"),
                compliance_score=int(audit_res.get("compliance_score", 100)),
                severity_level=audit_res.get("severity_level", "Low"),
                gap_details=audit_res.get("gap_details", ""),
                corrective_actions=audit_res.get("corrective_actions", [])
            )
            db.add(compliance_finding)
            logger.info(f"Generated compliance audit for {filename}: score={compliance_finding.compliance_score}")
        except Exception as e:
            logger.error(f"Failed compliance audit for {filename}: {e}")

    await db.commit()
    await db.refresh(doc)

    logger.info(f"Successfully ingested {filename} (id={doc.id}, chunks={doc.chunk_count})")
    return doc


async def ingest_text_document(
    text: str,
    filename: str,
    db: AsyncSession,
) -> Document:
    """Convenience wrapper: ingest a raw text string as if it were an uploaded file."""
    return await ingest_document(text.encode("utf-8"), filename, db)
